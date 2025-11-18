import re
import os
import json
import pdfplumber
import docx
import google.generativeai as genai
from typing import List, Dict, Union

# ──────────────────────────────────────────────
# CONFIGURATION
# ──────────────────────────────────────────────

# 1. Get key from Environment Variables (Preferred) or use the hardcoded key
API_KEY = os.environ.get("GEMINI_API_KEY")

# Fallback: Use the specific key provided if environment variable is missing
if not API_KEY:
    API_KEY = "AIzaSyDtsgMX4-ZUdbkMGuvv5uf3adD35iWBG2U"

if API_KEY:
    genai.configure(api_key=API_KEY)
    # Updated to the specific preview model requested
    model = genai.GenerativeModel("gemini-2.5-flash-preview-09-2025")
else:
    print("⚠️ WARNING: GEMINI_API_KEY not found.")
    print("AI features will be disabled. Please set the GEMINI_API_KEY environment variable.")
    model = None

# ──────────────────────────────────────────────
# 1. CLASSICAL REGEX EXTRACTORS (Fallback/Fast)
# ──────────────────────────────────────────────

def extract_email(text: str) -> str:
    # Improved regex to catch more email formats
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None

def extract_phone(text: str) -> str:
    # Supports multiple formats: +91 9876543210, 9876543210, 123-456-7890
    match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    return match.group(0) if match else None

def extract_experience(text: str) -> str:
    # Looks for patterns like "5+ years", "5 years", "5.5 years"
    match = re.search(r"(\d+(\.\d+)?)\+?\s+years?", text, re.I)
    return match.group(1) if match else "0"

def extract_location(text: str) -> str:
    # Common Indian IT hubs + major global cities
    cities = [
        "Mumbai", "Bangalore", "Bengaluru", "Delhi", "Pune",
        "Hyderabad", "Chennai", "Kolkata", "Noida", "Gurgaon",
        "San Francisco", "New York", "London", "Remote"
    ]
    for c in cities:
        # Word boundary check to avoid partial matches (e.g., 'Delhis' vs 'Delhi')
        if re.search(r"\b" + re.escape(c) + r"\b", text, re.IGNORECASE):
            return c
    return "Unknown"

def extract_job_title(text: str) -> str:
    titles = [
        "DevOps Engineer", "Software Engineer", "Cloud Engineer",
        "Python Developer", "Frontend Developer", "Backend Developer",
        "Full Stack Developer", "Data Scientist", "ML Engineer",
        "System Administrator", "Project Manager"
    ]
    for t in titles:
        if t.lower() in text.lower():
            return t
    return "Unknown"

def extract_skills_basic(text: str) -> List[str]:
    """Simple keyword matching for common skills."""
    skill_db = [
        "python", "java", "c++", "aws", "docker", "kubernetes", "devops",
        "react", "node", "sql", "git", "linux", "terraform",
        "jenkins", "mongodb", "azure", "gcp", "flask", "django", "html", "css"
    ]
    found = []
    for skill in skill_db:
        # Use regex word boundaries to ensure "C" doesn't match "Claude"
        if re.search(r"\b" + re.escape(skill) + r"\b", text, re.IGNORECASE):
            found.append(skill)
    return list(set(found))

# ──────────────────────────────────────────────
# 2. AI EXTRACTION – GEMINI POWERED
# ──────────────────────────────────────────────

def extract_skills_ai(text: str) -> List[str]:
    """
    Uses Gemini to identify skills more accurately and return valid JSON.
    """
    if not model:
        return []

    prompt = f"""
    Analyze the following resume text and extract technical skills.
    Return ONLY a raw JSON list of strings. Do not include Markdown formatting or backticks.
    
    Example output: ["Python", "Java", "Team Leadership"]
    
    Resume Text:
    {text[:3000]}  # Truncate to avoid token limits if text is massive
    """

    try:
        response = model.generate_content(prompt)
        
        # Clean the response to ensure it's valid JSON
        cleaned_text = response.text.strip()
        if cleaned_text.startswith("```json"):
            cleaned_text = cleaned_text[7:-3]
        elif cleaned_text.startswith("```"):
            cleaned_text = cleaned_text[3:-3]
            
        skills = json.loads(cleaned_text)
        
        # Ensure it is a list
        if isinstance(skills, list):
            return [str(s).strip() for s in skills]
        else:
            return []
            
    except Exception as e:
        print(f"Error in AI skill extraction: {e}")
        return []

def generate_resume_improvement(text: str) -> str:
    """
    Generates professional resume improvement suggestions.
    """
    if not model:
        return "AI suggestions unavailable (Gemini API key not set)."

    prompt = f"""
    You are a senior Technical Recruiter. Analyze this resume text for a developer role.
    Provide a structured critique in plain text:
    1. Key Strengths
    2. Missing Critical Skills (based on current market trends)
    3. Formatting & Clarity Issues
    4. A better 2-line Professional Summary
    
    Resume Text:
    {text[:4000]}
    """

    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Unable to generate AI suggestions: {e}"

# ──────────────────────────────────────────────
# 3. FILE PARSING UTILITIES
# ──────────────────────────────────────────────

def extract_text_from_pdf(file_source) -> str:
    text = ""
    try:
        # pdfplumber.open accepts both file paths and file-like objects
        with pdfplumber.open(file_source) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(file_source) -> str:
    try:
        # docx.Document accepts both file paths and file-like objects
        doc = docx.Document(file_source)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

# ──────────────────────────────────────────────
# 4. MAIN ORCHESTRATOR
# ──────────────────────────────────────────────

def parse_resume(file_input) -> Dict[str, Union[str, List[str]]]:
    """
    Main entry point to parse a resume file.
    Accepts either a file path (str) or a file-like object (from Streamlit/Flask).
    """
    # Determine if input is a string (file path) or an object (uploaded file)
    if isinstance(file_input, str):
        if not os.path.exists(file_input):
            return {"error": "File not found"}
        ext = file_input.split(".")[-1].lower()
    else:
        # Assume it's a file object (like Streamlit's UploadedFile)
        # These objects usually have a 'name' attribute
        if hasattr(file_input, 'name'):
             ext = file_input.name.split(".")[-1].lower()
        else:
             return {"error": "Unknown file object format"}

    # Step 1: Extract Raw Text
    # Note: We pass the file_input directly to the extractors. 
    # Libraries like pdfplumber and python-docx handle file-like objects automatically.
    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_input)
    elif ext == "docx":
        raw_text = extract_text_from_docx(file_input)
    else:
        return {"error": "Unsupported file format. Please use PDF or DOCX."}

    if not raw_text.strip():
        return {"error": "Could not extract text from file. It might be an image-based PDF."}

    # Step 2: Classical Regex Extraction (Fast)
    print("running regex extraction...")
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    basic_skills = extract_skills_basic(raw_text)
    exp = extract_experience(raw_text)
    loc = extract_location(raw_text)
    title = extract_job_title(raw_text)

    # Step 3: AI Analysis (Smart)
    print("running Gemini AI extraction...")
    ai_skills = extract_skills_ai(raw_text)
    suggestions = generate_resume_improvement(raw_text)

    # Merge skills: prioritized AI skills, but keep regex ones if AI missed them
    final_skills = list(set(basic_skills + ai_skills))

    return {
        "candidate_info": {
            "email": email,
            "phone": phone,
            "location": loc,
            "experience_years": exp,
            "likely_job_title": title,
        },
        "skills": final_skills,
        "ai_suggestions": suggestions,
        "raw_text_snippet": raw_text[:500] + "..." # Preview
    }

# ──────────────────────────────────────────────
# EXECUTION BLOCK (TESTING)
# ──────────────────────────────────────────────

if __name__ == "__main__":
    # 1. Create a dummy resume for testing if no file exists
    dummy_filename = "test_resume.docx"
    
    if not os.path.exists(dummy_filename):
        print(f"Creating dummy file '{dummy_filename}' for testing...")
        doc = docx.Document()
        doc.add_heading('Dev Singh', 0)
        doc.add_paragraph('Email: dev.singh@mca-student.com | Phone: +91 9876543210')
        doc.add_paragraph('Location: Mumbai, India')
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('Junior Software Engineer - 2 Years Experience')
        doc.add_paragraph('Worked on Python, Django, and React projects.')
        doc.save(dummy_filename)

    # 2. Run the parser
    result = parse_resume(dummy_filename)

    # 3. Print Results
    print("\n" + "="*40)
    print(" RESUME PARSING RESULT ")
    print("="*40)
    
    if "error" in result:
        print(f"Error: {result['error']}")
    else:
        print(f"📧 Email: {result['candidate_info']['email']}")
        print(f"📱 Phone: {result['candidate_info']['phone']}")
        print(f"📍 Location: {result['candidate_info']['location']}")
        print(f"💼 Experience: {result['candidate_info']['experience_years']} years")
        print(f"🛠️ Skills Found: {', '.join(result['skills'])}")
        print("-" * 40)
        print("🤖 AI SUGGESTIONS:")
        print(result['ai_suggestions'])